import requests
from docx import Document
import os

def create_doc(details: dict, filename: str):
    doc = Document()
    for key, value in details.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(filename)

def read_doc_from_cid(cid: str):
    """
    Downloads a .docx file from IPFS using the given CID and reads its content as plain text.
    """
    try:
        url = f"https://gateway.pinata.cloud/ipfs/{cid}"
        response = requests.get(url)
        if response.status_code != 200:
            return f"❌ Failed to fetch file. Status code: {response.status_code}"

        with open("temp.docx", "wb") as f:
            f.write(response.content)

        doc = Document("temp.docx")
        text = "\n".join([p.text for p in doc.paragraphs])
        os.remove("temp.docx")
        return text
    except Exception as e:
        return f"❌ Error reading document: {str(e)}"

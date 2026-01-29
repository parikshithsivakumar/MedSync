import gradio as gr
from config import contract, w3
from word_utils import create_doc
from pinata_utils import upload_to_pinata
from web3.exceptions import ContractLogicError
from eth_utils import is_address
from docx import Document
import os
import requests

custom_css = """
    @import url('https://fonts.googleapis.com/css2?family=Lato:wght@300;400;700&display=swap');

    /* Base body style */
    body {
        background: linear-gradient(135deg, #fdfbfb 0%, #ebedee 100%);
        font-family: 'Lato', sans-serif;
        margin: 0;
        padding-top: 30px;
        overflow-x: hidden;
    }

    /* Header box without transform */
    #header-box {
        background-color: #ffffffee;
        border-radius: 16px;
        padding: 28px;
        margin-bottom: 28px;
        box-shadow: 0 6px 16px rgba(0, 0, 0, 0.06);
        border: 3px solid #ED4C97;
    }
    #header-box:hover, #register-box:hover, #upload-box:hover,
    #doctor-dashboard-box:hover, #patient-view-box:hover, #lab-tech-view-box:hover {
        box-shadow: 0 12px 24px rgba(0, 0, 0, 0.08);
    }

    /* Card boxes */
    #register-box, #upload-box, #doctor-dashboard-box, #patient-view-box, #lab-tech-view-box {
        background-color: white;
        border-radius: 12px;
        padding: 30px;
        border: 3px solid #f06292;
        box-shadow: 0 4px 12px rgba(240, 98, 146, 0.2);
        margin-bottom: 30px;
        transition: box-shadow 0.3s ease, border-color 0.3s ease;
    }

    /* Tabs */
    .gr-tab-nav {
        background-color: transparent;
        border-bottom: 2px solid #e0e0e0;
    }
    .gr-tab-nav .gr-tab {
        padding: 10px 16px;
        color: #9c27b0;
        font-weight: 600;
        border-bottom: 3px solid transparent;
        transition: all 0.2s ease-in-out;
    }
    .gr-tab-nav .gr-tab:hover {
        color: #6a1b9a;
    }
    .gr-tab-nav .gr-tab.gr-tab-selected {
    color: #9e9e9e; /* Gray color */
    border-bottom: 3px solid #9e9e9e; /* Gray border for the active tab */
}


    .file-dropbox {
        border: 2px dashed #f06292 !important;
        background-color: #fff0f6 !important;
        border-radius: 8px;
    }

    button {
        background-color: #ec407a !important;
        color: white !important;
        font-weight: bold !important;
        font-size: 16px !important;
        border-radius: 10px !important;
        padding: 14px 28px !important;
        margin: 10px 8px !important;
        transition: background 0.3s ease, transform 0.2s ease;
        font-family: 'lato', cursive !important;
    }
    button:hover {
        background-color: #d81b60 !important;
        transform: scale(1.03);
    }

    textarea, input, select {
        border: 1px solid #e0e0e0 !important;
        border-radius: 6px !important;
        padding: 10px !important;
        font-size: 14px !important;
        width: 100%;
        font-family: 'lato', cursive !important;
    }

    label {
        font-weight: 600;
        color: #333;
        margin-bottom: 6px;
        display: block;
        font-family: 'lato', cursive !important;
    }

    .gr-dropdown {
        position: relative !important;
        z-index: 1000 !important;
    }

    .gr-dropdown .gr-select {
        position: relative !important;
        z-index: 1010 !important;
    }

    .gr-dropdown .gr-select > div[role="listbox"] {
        position: absolute !important;
        top: 100% !important;
        left: 0 !important;
        right: 0 !important;
        margin-top: 4px;
        z-index: 2000 !important;
        background-color: white;
        border-radius: 6px;
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
        max-height: 200px;
        overflow-y: auto;
    }

    textarea:focus, input:focus {
        border-color: #ff2e91 !important;
        box-shadow: 0 0 12px 2px rgba(255, 46, 145, 0.7);
        outline: none;
    }

    select:focus {
        border-color: #ff2e91 !important;
        box-shadow: 0 0 12px 2px rgba(255, 46, 145, 0.7);
        outline: none;
    }

    input, textarea {
        max-width: 100% !important;
        box-sizing: border-box;
    }

    #doctor-dashboard-box {
        border: 3px solid #f06292;
    }
"""


def load_doc_from_ipfs(cid):
    url = f"https://gateway.pinata.cloud/ipfs/{cid}"
    response = requests.get(url)
    with open("temp.docx", "wb") as f:
        f.write(response.content)
    doc = Document("temp.docx")
    text = "\n".join([p.text for p in doc.paragraphs])
    os.remove("temp.docx")
    return text

def register(role, wallet_address, name):
    wallet_address = wallet_address.strip()
    if not is_address(wallet_address):
        return "❌ Invalid wallet address."
    try:
        role_map = {"Patient": 1, "Doctor": 2, "Lab Technician": 3}
        tx = contract.functions.register(role_map[role], name).transact({'from': wallet_address})
        w3.eth.wait_for_transaction_receipt(tx)
        return f"✅ {name} ({role}) registered successfully!"
    except ContractLogicError as e:
        return f"❌ Registration failed: {str(e)}"
    except Exception as e:
        return f"❌ Error: {str(e)}"

def get_user_name(wallet_address):
    try:
        name = contract.functions.getName(wallet_address).call()
        return name
    except Exception as e:
        return f"❌ Error: {str(e)}"

def doctor_dashboard(wallet_address, patient_address, name, age, diagnosis, treatment, scan, approval):
    wallet_address = wallet_address.strip()
    patient_address = patient_address.strip()
    if not is_address(wallet_address) or not is_address(patient_address):
        return "❌ Invalid wallet or patient address."
    try:
        details = {"Name": name,"Age": age,"Diagnosis": diagnosis,"Treatment": treatment,"Scan": scan,"Scan Approved for Lab": approval}
        filename = f"{patient_address}.docx"
        create_doc(details, filename)
        cid = upload_to_pinata(filename)
        tx = contract.functions.uploadFile(patient_address, cid).transact({'from': wallet_address})
        w3.eth.wait_for_transaction_receipt(tx)
        return f"✅ Report uploaded to IPFS. CID: {cid}"
    except ContractLogicError:
        return "❌ Upload failed: Only doctors can upload."
    except Exception as e:
        return f"❌ Error uploading: {str(e)}"

def edit_doc(content, wallet_address, patient_address):
    if not is_address(patient_address):
        return "❌ Invalid patient address."
    try:
        filename = f"{patient_address}_edited.docx"
        create_doc({"Edited Report": content}, filename)
        cid = upload_to_pinata(filename)
        tx = contract.functions.uploadFile(patient_address, cid).transact({'from': wallet_address})
        w3.eth.wait_for_transaction_receipt(tx)
        return f"✅ Updated report uploaded. CID: {cid}"
    except Exception as e:
        return f"❌ Error during edit: {str(e)}"

def lab_tech_view(wallet_address, patient_address):
    if not is_address(wallet_address) or not is_address(patient_address):
        return "❌ Invalid wallet or patient address."
    try:
        cids = contract.functions.getFiles(patient_address).call()
        scan_cids = []
        for cid in cids:
            try:
                response = requests.get(f"https://gateway.pinata.cloud/ipfs/{cid}")
                with open("temp_lab.docx", "wb") as f:
                    f.write(response.content)
                doc = Document("temp_lab.docx")
                os.remove("temp_lab.docx")
                if any("Scan Approved for Lab: Yes" in p.text for p in doc.paragraphs):
                    scan_cids.append(cid)
            except:
                continue
        if not scan_cids:
            return "No scan-related files found."
        links = [f'<a href="https://gateway.pinata.cloud/ipfs/{cid}" target="_blank">📄 View Scan Document ({cid[:8]}...)</a>' for cid in scan_cids]
        return "\n".join(links)
    except Exception as e:
        return f"❌ Error: {str(e)}"

def upload_scan_file(wallet_address, patient_address, scan_type, scan_file):
    try:
        cid = upload_to_pinata(scan_file.name)
        filename = f"{patient_address}_{scan_type}.docx"
        create_doc({"Scan Type": scan_type, "Scan Image IPFS": f"https://gateway.pinata.cloud/ipfs/{cid}"}, filename)
        doc_cid = upload_to_pinata(filename)
        tx = contract.functions.uploadFile(patient_address, doc_cid).transact({'from': wallet_address})
        w3.eth.wait_for_transaction_receipt(tx)
        return f"✅ {scan_type} scan uploaded. CID: {doc_cid}"
    except Exception as e:
        return f"✅ {scan_type} scan uploaded. CID: {doc_cid}"


def patient_view(wallet_address):
    wallet_address = wallet_address.strip()
    if not is_address(wallet_address):
        return "<p style='color:pink;'>❌ Invalid wallet address.</p>"
    try:
        cids = contract.functions.getFiles(wallet_address).call()
        if not cids:
            return "<p>No records found.</p>"
        links_html = "".join([f"<p>File {i+1}: <a href='https://gateway.pinata.cloud/ipfs/{cid}' target='_blank'>Open File</a></p>" for i, cid in enumerate(cids)])
        return links_html
    except Exception as e:
        return f"<p style='color:pink;'>❌ Error: {str(e)}</p>"

with gr.Blocks(css=custom_css) as demo:
    gr.Markdown("""
<link href="https://fonts.googleapis.com/css2?family=Sofia&display=swap" rel="stylesheet">
<div id='header-box' style='display:flex;align-items:center;'>
    <span style='font-size:2.2em; margin-right:15px; color:#ED4C97;'>❤️</span>
    <div>
        <span style='font-size:2.2em; color:#ED4C97; font-weight:bold; font-family: "lato", cursive;'>MedSync</span><br>
        <span style='font-size:1.2em; color:#ED4C97;'>Secure Medical Records</span>
    </div>
</div>
""")


    with gr.Tab("Register"):
        with gr.Column(elem_id="register-box"):
            name_input = gr.Textbox(label="Name")
            role = gr.Dropdown(["Patient", "Doctor", "Lab Technician"], label="Role")
            wallet = gr.Textbox(label="Wallet Address")
            out = gr.Textbox()
            btn = gr.Button("Register")
            btn.click(fn=register, inputs=[role, wallet, name_input], outputs=out)

    with gr.Tab("Doctor Dashboard"):
        with gr.Column(elem_id="doctor-dashboard-box"):
            doc_wallet = gr.Textbox(label="Doctor Wallet Address")
            patient_wallet = gr.Textbox(label="Patient Wallet Address")
            name = gr.Textbox(label="Patient Name")
            age = gr.Textbox(label="Age")
            diagnosis = gr.Textbox(label="Diagnosis")
            treatment = gr.Textbox(label="Treatment")
            scan = gr.Textbox(label="Scan Result")
            approval = gr.Dropdown(label="Allow Lab Technician Access to Scan?", choices=["Yes", "No"])
            doc_out = gr.Textbox(label="Status")
            doc_btn = gr.Button("Upload Report")
            doc_btn.click(fn=doctor_dashboard, inputs=[doc_wallet, patient_wallet, name, age, diagnosis, treatment, scan, approval], outputs=doc_out)

            last_cid = gr.Textbox(label="Last File CID")
            fetch_btn = gr.Button("Fetch Latest CID")
            fetch_btn.click(fn=lambda patient_address: contract.functions.getFiles(patient_address).call()[-1] if contract.functions.getFiles(patient_address).call() else "No records", inputs=[patient_wallet], outputs=last_cid)

            view_btn = gr.Button("View Latest Report")
            doc_text = gr.Textbox(label="View/Edit Document", lines=15)
            view_btn.click(fn=load_doc_from_ipfs, inputs=last_cid, outputs=doc_text)

            edit_btn = gr.Button("Save & Upload Edited Report")
            edit_output = gr.Textbox(label="Edit Upload Status")
            edit_btn.click(fn=edit_doc, inputs=[doc_text, doc_wallet, patient_wallet], outputs=edit_output)

    with gr.Tab("Lab Technician View"):
        with gr.Column(elem_id="lab-tech-view-box"):
            lab_wallet = gr.Textbox(label="Lab Technician Wallet Address")
            name = gr.Textbox(label="Patient Name")
            patient_address = gr.Textbox(label="Patient Wallet Address")
            lab_out = gr.HTML(label="Files with Scan")
            lab_btn = gr.Button("View Files")
            lab_btn.click(fn=lab_tech_view, inputs=[lab_wallet, patient_address], outputs=lab_out)

            scan_type = gr.Dropdown(label="Scan Type", choices=["X-ray", "CT scan", "MRI scan", "Ultrasound", "PET scan"])
            scan_file = gr.File(label="Upload Scan Image")
            upload_out = gr.Textbox(label="Upload Status")
            upload_btn = gr.Button("Upload Scan Image")
            upload_btn.click(fn=upload_scan_file, inputs=[lab_wallet, patient_address, scan_type, scan_file], outputs=upload_out)

            view_btn = gr.Button("View Latest Report")
            doc_text = gr.Textbox(label="View/Edit Document", lines=15)
            view_btn.click(fn=load_doc_from_ipfs, inputs=last_cid, outputs=doc_text)

            edit_btn = gr.Button("Save & Upload Edited Report")
            edit_output = gr.Textbox(label="Edit Upload Status")
            edit_btn.click(fn=edit_doc, inputs=[doc_text, doc_wallet, patient_wallet], outputs=edit_output)

    with gr.Tab("Patient View"):
         with gr.Column(elem_id="patient-view-box"):
            patient_wallet_view = gr.Textbox(label="Your Wallet Address")
            name = gr.Textbox(label="Your Name")
            view_out = gr.HTML(label="Your Files")
            view_btn = gr.Button("View My Files")
            view_btn.click(fn=patient_view, inputs=[patient_wallet_view], outputs=view_out)

if __name__ == "__main__":
    demo.launch()

